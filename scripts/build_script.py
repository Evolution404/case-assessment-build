#!/usr/bin/env python3
import json
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CFG = json.loads((ROOT / "content/case.json").read_text(encoding="utf-8"))
DEFENSE = json.loads((ROOT / "content/defense.json").read_text(encoding="utf-8"))
OUT = ROOT / "dist/答辩逐字稿-从人海作业到数智协同.docx"
CASE_ID = os.environ.get("CASE_ID", CFG["case_id_default"])

# compact_reference_guide + named rehearsal overrides
LATIN_FONT = "Arial"
CN_FONT = "Hiragino Sans GB"
BODY_SIZE = 15
BODY_LINE = 25
RUST = "BD6549"
BLUE = "537484"
INK = "1D1C1A"
MUTED = "756F67"
PAPER = "F4F0E9"


def speaker_tokens():
    m = CFG["metrics"]
    return {
        "province_poles_wan1": f"{m['province_poles'] / 10000:.1f}万",
        "province_lines_label": m["province_lines_report_label"],
        "alarm_photos_comma": f"{m['alarm_photos']:,}",
        "patrol_photos_wan0": f"{round(m['patrol_photos_monthly'] / 10000)}万",
        "crossing_poles_comma": f"{m['crossing_poles']:,}",
        "crossing_lines": str(m["crossing_lines"]),
        "crossing_records": str(m["crossing_existing_records"]),
        "crossing_people": m["crossing_people_before"],
        "crossing_before": m["crossing_duration_before"],
        "crossing_after": m["crossing_duration_after"],
        "crossing_findings": m["crossing_additional_findings"],
        "fireworks_turnaround": m["fireworks_turnaround"],
        "alarm_candidates_comma": f"{m['alarm_candidates']:,}",
        "alarm_reviewed_comma": f"{m['alarm_reviewed']:,}",
        "alarm_confirmed_comma": f"{m['alarm_confirmed_pairs']:,}",
        "alarm_different_comma": f"{m['alarm_confirmed_different']:,}",
        "alarm_pending_comma": f"{m['alarm_pending']:,}",
        "alarm_review_rate": f"{m['alarm_candidate_hit_rate']:.2f}%",
        "trial_total": str(m["province_trial_duplicate_issues"]),
        "trial_alarm": str(m["province_trial_alarm_duplicate_issues"]),
        "trial_patrol": str(m["province_trial_patrol_duplicate_issues"]),
        "theoretical_pairs_wan_yi": f"{m['theoretical_pairs'] / 1e12:.1f}万亿",
        "patrol_pairs_comma": f"{m['patrol_duplicate_pairs']:,}",
        "pairs_per_duplicate_yi": f"{round(m['pairs_per_duplicate'] / 1e8)}亿",
    }


TOKENS = speaker_tokens()


def resolve(text):
    value = re.sub(r"\{\{([a-z0-9_]+)\}\}", lambda m: TOKENS.get(m.group(1), m.group(0)), text)
    if "{{" in value:
        raise ValueError(f"未解析的逐字稿变量：{value}")
    return value


def set_run(run, size, bold=False, color=INK, font=LATIN_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    fonts.set(qn("w:eastAsia"), CN_FONT)
    fonts.set(qn("w:cs"), font)


def add_paragraph(doc, text, size=BODY_SIZE, bold=False, color=INK, align=None,
                  before=0, after=6, line=BODY_LINE, first_indent=False):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    if first_indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        para.alignment = align
    set_run(para.add_run(text), size, bold, color)
    return para


def add_left_rule(para, color=RUST, size="16", space="130"):
    ppr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    borders.append(left)
    ppr.append(borders)


def shade(para, fill=PAPER):
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_page_field(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(para.add_run("第 "), 9, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    para._p.append(fld)
    set_run(para.add_run(" / 12 页"), 9, False, MUTED)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def setup_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("个人案例答辩｜从人海作业到数智协同"), 9, True, MUTED)
    add_page_field(section.footer.paragraphs[0])


def build():
    doc = Document()
    setup_styles(doc)
    setup_section(doc.sections[0])
    doc.core_properties.title = "答辩逐字稿｜从人海作业到数智协同"
    doc.core_properties.subject = "12分钟个人案例答辩逐字稿"
    doc.core_properties.author = "案例答辩"
    doc.core_properties.keywords = "外协管理, 增效, 提质, 数智协同"

    slides = DEFENSE["slides"]
    for index, item in enumerate(slides):
        if index == 0:
            add_paragraph(doc, "答辩逐字稿", 25, True, INK, WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, line=34)
            add_paragraph(doc, CFG["title"] + "——" + CFG["subtitle"], 13.5, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=3, line=23)
            add_paragraph(doc, CASE_ID, 10.5, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=18, line=18)
        else:
            add_paragraph(doc, "个人案例答辩", 9.5, True, RUST, before=4, after=12, line=16)

        title_text = f"第{item['no']}页｜{item['title']}"
        title = add_paragraph(doc, title_text, 19, True, INK, before=2, after=8, line=29)
        add_left_rule(title)

        meta = add_paragraph(
            doc,
            f"时间：{item['time']}\n操作：{item['cue']}\n过渡：{item['transition']}",
            10.5,
            False,
            MUTED,
            before=2,
            after=16,
            line=18,
        )
        meta.paragraph_format.left_indent = Pt(12)
        meta.paragraph_format.right_indent = Pt(12)
        shade(meta, "F1E9E1")

        add_paragraph(doc, resolve(item["script"]), BODY_SIZE, False, INK, before=2, after=10, line=BODY_LINE, first_indent=True)

        if item.get("sources"):
            add_paragraph(doc, "资料口径：" + "；".join(item["sources"]), 9.5, False, MUTED, before=12, after=0, line=16)

        if index != len(slides) - 1:
            doc.add_page_break()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    total_chars = sum(len(resolve(item["script"])) for item in slides)
    print(f"[script] 12 pages, {total_chars} Chinese-script characters -> {OUT}")


if __name__ == "__main__":
    build()
