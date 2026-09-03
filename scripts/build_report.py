#!/usr/bin/env python3
import json, tempfile
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
TEMPLATE=ROOT.parent/"06 研讨报告模板【研讨报告可参考，根据实际.docx"
OUT=ROOT/"dist/案例考核报告-从人海作业到数智协同.docx"
CONFIRM_FIG=ROOT/"dist/confirm"
CFG=json.loads((ROOT/"content/case.json").read_text(encoding="utf-8"))
ALARM_SCOPE=CFG["metrics"]["alarm_scope_label"]
PROVINCE_LINES_LABEL=CFG["metrics"]["province_lines_report_label"]
FONT="仿宋_GB2312"
DOCX_MEDIA_DIR=Path(tempfile.gettempdir())/"case-assessment-docx-media"

def wipe_body(doc):
    body=doc._element.body
    for child in list(body):
        if child.tag!=qn("w:sectPr"):
            body.remove(child)

def set_run(run,size=14,bold=False,color="000000",font=FONT):
    run.font.name=font;run.font.size=Pt(size);run.font.bold=bold;run.font.color.rgb=RGBColor.from_string(color)
    rpr=run._element.get_or_add_rPr();rf=rpr.find(qn("w:rFonts"))
    if rf is None: rf=OxmlElement("w:rFonts");rpr.append(rf)
    for k in ("ascii","hAnsi","eastAsia","cs"):rf.set(qn("w:"+k),font)

def para(doc,text="",size=14,bold=False,align=None,indent=True,before=0,after=0,line=28,color="000000",keep=False,together=False):
    p=doc.add_paragraph();f=p.paragraph_format
    f.line_spacing_rule=WD_LINE_SPACING.EXACTLY;f.line_spacing=Pt(line);f.space_before=Pt(before);f.space_after=Pt(after)
    f.widow_control=True
    if align is not None:f.alignment=align
    if indent:f.first_line_indent=Pt(size*2)
    if keep:f.keep_with_next=True
    if together:f.keep_together=True
    set_run(p.add_run(text),size,bold,color)
    return p

def _set_outline_level(p,level):
    ppr=p._p.get_or_add_pPr();outline=ppr.find(qn("w:outlineLvl"))
    if outline is None:outline=OxmlElement("w:outlineLvl");ppr.append(outline)
    outline.set(qn("w:val"),str(level))

def h1(doc,text):
    p=para(doc,text,14,True,indent=True,keep=True);_set_outline_level(p,0);return p

def h2(doc,text):
    p=para(doc,text,14,True,indent=True,keep=True)
    _set_outline_level(p,2 if text[:1].isdigit() else 1)
    return p

def bullet(doc,text):return para(doc,text,14,False,indent=True)

def report_title(doc,title,subtitle):
    p=para(doc,title,18,False,WD_ALIGN_PARAGRAPH.CENTER,False,line=28)
    p=doc.add_paragraph();f=p.paragraph_format
    f.line_spacing_rule=WD_LINE_SPACING.EXACTLY;f.line_spacing=Pt(28);p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("—"),16,False)
    set_run(p.add_run(subtitle),16,False)
    return p

def prepare_docx_image(image_path,width_cm):
    """Resize + JPEG-compress figures only for DOCX embedding; source figures stay untouched."""
    DOCX_MEDIA_DIR.mkdir(parents=True,exist_ok=True)
    target_w=max(900,int(width_cm/2.54*220))
    out=DOCX_MEDIA_DIR/(image_path.stem+f"-{target_w}px.jpg")
    src_mtime=image_path.stat().st_mtime_ns
    stamp=DOCX_MEDIA_DIR/(out.name+".mtime")
    if out.exists() and stamp.exists() and stamp.read_text()==str(src_mtime):
        return out
    with Image.open(image_path) as im:
        im=im.convert("RGBA")
        bg=Image.new("RGB",im.size,"white");bg.paste(im,mask=im.getchannel("A"))
        if bg.width>target_w:
            target_h=max(1,round(bg.height*target_w/bg.width))
            bg=bg.resize((target_w,target_h),Image.Resampling.LANCZOS)
        bg.save(out,"JPEG",quality=90,optimize=True,progressive=True,subsampling=0,dpi=(220,220))
    stamp.write_text(str(src_mtime))
    return out

def figure(doc,name,caption,width=15.0):
    confirmed={
        "pain-points":"01-外协管理两大痛点.png",
        "province-map":"02-省域线路杆塔任务规模.png",
        "dual-wheel":"03-增效提质总体模型.png",
        "workflow":"04-外协任务数字化筛选流程.png",
        "crossing":"05-交叉跨越自动筛查.png",
        "bird":"06-鸟类活动重点区域筛查.png",
        "fireworks":"07-集中燃放点周边杆塔筛查.png",
        "photo-scale":"08-告警工单照片筛选-1-核心规模与筛选复核.png",
        "photo-cases":"08-告警工单照片筛选-2-真实案例.png",
        "outcomes":"10-外协管理模式转型总结.png",
    }
    if name not in confirmed:
        raise KeyError(f"未配置的确认插图：{name}")
    image_path=CONFIRM_FIG/confirmed[name]
    if not image_path.exists():
        raise FileNotFoundError(f"缺少确认插图：{image_path}")
    embedded_image=prepare_docx_image(image_path,width)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(5);p.paragraph_format.space_after=Pt(0);p.paragraph_format.keep_with_next=True
    shape=p.add_run().add_picture(str(embedded_image),width=Cm(width))
    shape._inline.docPr.set("title",caption)
    shape._inline.docPr.set("descr",caption)
    c=doc.add_paragraph();c.alignment=WD_ALIGN_PARAGRAPH.CENTER;c.paragraph_format.space_after=Pt(0);c.paragraph_format.first_line_indent=Pt(0)
    c.paragraph_format.line_spacing_rule=WD_LINE_SPACING.EXACTLY;c.paragraph_format.line_spacing=Pt(24)
    set_run(c.add_run(caption.replace("\u3000","  ")),12,False,"000000")

def page_break(doc):
    # 模板采用连续研讨报告排版，不主动插入分页符；由 Word 根据正文与图片自然分页。
    return None

def _set_table_edge(parent,edge,val="single",sz="8"):
    borders=parent.find(qn("w:tblBorders"))
    if borders is None:
        borders=OxmlElement("w:tblBorders");parent.append(borders)
    node=borders.find(qn("w:"+edge))
    if node is None:
        node=OxmlElement("w:"+edge);borders.append(node)
    node.set(qn("w:val"),val);node.set(qn("w:sz"),sz);node.set(qn("w:space"),"0");node.set(qn("w:color"),"000000")

def _set_cell_bottom_border(cell,sz="6"):
    tcpr=cell._tc.get_or_add_tcPr();borders=tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders=OxmlElement("w:tcBorders");tcpr.append(borders)
    bottom=borders.find(qn("w:bottom"))
    if bottom is None:
        bottom=OxmlElement("w:bottom");borders.append(bottom)
    bottom.set(qn("w:val"),"single");bottom.set(qn("w:sz"),sz);bottom.set(qn("w:space"),"0");bottom.set(qn("w:color"),"000000")

def paper_table(doc,caption,headers,rows,widths=None,note=None):
    """Formal three-line table for the report: no fill, no vertical rules, compact academic typesetting."""
    cap=doc.add_paragraph();cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent=Pt(0);cap.paragraph_format.space_before=Pt(6);cap.paragraph_format.space_after=Pt(4)
    cap.paragraph_format.keep_with_next=True
    set_run(cap.add_run(caption.replace("\u3000","  ")),11,False,"000000")
    tbl=doc.add_table(rows=1,cols=len(headers));tbl.alignment=WD_TABLE_ALIGNMENT.CENTER;tbl.autofit=False
    tblpr=tbl._tbl.tblPr
    for edge in ("left","right","insideH","insideV"):_set_table_edge(tblpr,edge,"nil","0")
    _set_table_edge(tblpr,"top","single","12");_set_table_edge(tblpr,"bottom","single","12")
    if widths is not None:
        # Fix both tblGrid and cell widths so Word/LibreOffice use the same column geometry.
        for ci,width in enumerate(widths):
            tbl.columns[ci].width=Cm(width)
            tbl._tbl.tblGrid.gridCol_lst[ci].set(qn("w:w"),str(Cm(width).twips))
        for row in tbl.rows:
            for ci,cell in enumerate(row.cells):cell.width=Cm(widths[ci])
    trpr=tbl.rows[0]._tr.get_or_add_trPr();header_repeat=OxmlElement("w:tblHeader");header_repeat.set(qn("w:val"),"true");trpr.append(header_repeat)
    for ci,text in enumerate(headers):
        cell=tbl.rows[0].cells[ci];cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths is not None:cell.width=Cm(widths[ci])
        p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.first_line_indent=Pt(0);p.paragraph_format.space_before=Pt(2);p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.EXACTLY;p.paragraph_format.line_spacing=Pt(18)
        set_run(p.add_run(text),10.5,True,"000000");_set_cell_bottom_border(cell,"6")
    for row_values in rows:
        row=tbl.add_row();row_pr=row._tr.get_or_add_trPr();cant_split=OxmlElement("w:cantSplit");row_pr.append(cant_split)
        for ci,text in enumerate(row_values):
            cell=row.cells[ci];cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths is not None:cell.width=Cm(widths[ci])
            p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Pt(0);p.paragraph_format.space_before=Pt(1);p.paragraph_format.space_after=Pt(1)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.EXACTLY;p.paragraph_format.line_spacing=Pt(18)
            set_run(p.add_run(str(text)),10.5,False,"000000")
    # These report tables are short; keep their rows as one visual unit whenever a page has room.
    # Word may still break a table that is taller than the available page, with the repeated header retained.
    for keep_row in tbl.rows[:-1]:
        for cell in keep_row.cells:
            for keep_p in cell.paragraphs:keep_p.paragraph_format.keep_with_next=True
    if note:
        # Keep the last data row together with its note so a page never begins with an orphaned table note.
        for cell in tbl.rows[-1].cells:
            for last_p in cell.paragraphs:last_p.paragraph_format.keep_with_next=True
        p=doc.add_paragraph();p.paragraph_format.first_line_indent=Pt(0);p.paragraph_format.space_before=Pt(2);p.paragraph_format.space_after=Pt(5)
        p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.EXACTLY;p.paragraph_format.line_spacing=Pt(17)
        set_run(p.add_run("注："+note),9.5,False,"444444")
    return tbl

def _m_run(text,plain=True):
    """Create an upright OMML math run for the report's engineering-paper formula style."""
    run=OxmlElement("m:r")
    if plain:
        rpr=OxmlElement("m:rPr")
        normal=OxmlElement("m:nor")
        rpr.append(normal)
        run.append(rpr)
    # Also force the underlying Word run to upright Cambria Math.  Word honors m:sty='p';
    # the explicit w:i=0 keeps the same result in alternate DOCX renderers.
    wrpr=OxmlElement("w:rPr")
    fonts=OxmlElement("w:rFonts")
    for key in ("ascii","hAnsi","eastAsia","cs"):fonts.set(qn("w:"+key),"Cambria Math")
    wrpr.append(fonts)
    italic=OxmlElement("w:i");italic.set(qn("w:val"),"0");wrpr.append(italic)
    italic_cs=OxmlElement("w:iCs");italic_cs.set(qn("w:val"),"0");wrpr.append(italic_cs)
    run.append(wrpr)
    node=OxmlElement("m:t");node.text=text;run.append(node)
    return run

def _m_sub(base,sub,base_plain=True):
    node=OxmlElement("m:sSub");node.append(OxmlElement("m:sSubPr"))
    expr=OxmlElement("m:e");expr.append(_m_run(base,base_plain));node.append(expr)
    sub_node=OxmlElement("m:sub");sub_node.append(_m_run(sub));node.append(sub_node)
    return node

def _m_subsup(base,sub,sup,base_plain=True):
    node=OxmlElement("m:sSubSup");node.append(OxmlElement("m:sSubSupPr"))
    expr=OxmlElement("m:e");expr.append(_m_run(base,base_plain));node.append(expr)
    sub_node=OxmlElement("m:sub");sub_node.append(_m_run(sub));node.append(sub_node)
    sup_node=OxmlElement("m:sup");sup_node.append(_m_run(sup));node.append(sup_node)
    return node

def _m_fraction(numerator_nodes,denominator_nodes):
    node=OxmlElement("m:f");node.append(OxmlElement("m:fPr"))
    num=OxmlElement("m:num");den=OxmlElement("m:den")
    for item in numerator_nodes:num.append(item)
    for item in denominator_nodes:den.append(item)
    node.append(num);node.append(den)
    return node

def _m_sum(sub_text,sup_text,body_nodes):
    node=OxmlElement("m:nary")
    pr=OxmlElement("m:naryPr");char=OxmlElement("m:chr");char.set(qn("m:val"),"∑")
    lim=OxmlElement("m:limLoc");lim.set(qn("m:val"),"undOvr");pr.append(char);pr.append(lim);node.append(pr)
    sub=OxmlElement("m:sub");sub.append(_m_run(sub_text));node.append(sub)
    sup=OxmlElement("m:sup");sup.append(_m_run(sup_text));node.append(sup)
    expr=OxmlElement("m:e")
    for item in body_nodes:expr.append(item)
    node.append(expr)
    return node

def _native_equation(kind):
    math=OxmlElement("m:oMath")
    if kind=="intersection":
        for item in (_m_run("P(t) = A + t(B − A),    "),_m_run("Q(u) = C + u(D − C),    "),_m_run("0 ≤ t,u ≤ 1")):math.append(item)
    elif kind=="bird_surface":
        math.append(_m_run("S = "))
        math.append(_m_subsup("M","3×3","7"))
        math.append(_m_run("["));math.append(_m_run("ln",True));math.append(_m_run("(1 + "))
        math.append(_m_sub("n","ij"));math.append(_m_run(")]"))
    elif kind=="hamming":
        math.append(_m_sub("d","H"));math.append(_m_run("("));math.append(_m_sub("h","1"));math.append(_m_run(","));math.append(_m_sub("h","2"));math.append(_m_run(") = "))
        body=[_m_run("I",True),_m_run("("),_m_sub("h","1i"),_m_run(" ≠ "),_m_sub("h","2i"),_m_run(")")]
        math.append(_m_sum("i=1","L",body))
    elif kind=="combinations":
        math.append(_m_run("C(N,2) = "))
        math.append(_m_fraction([_m_run("N(N − 1)")],[_m_run("2")]))
    else:
        raise KeyError(f"未知公式类型：{kind}")
    return math

def paper_equation(doc,kind,number):
    """Insert an upright, native editable Word OMML equation with a right-aligned equation number."""
    p=doc.add_paragraph();f=p.paragraph_format
    f.first_line_indent=Pt(0);f.space_before=Pt(6);f.space_after=Pt(6);f.keep_with_next=True;f.keep_together=True
    f.line_spacing_rule=WD_LINE_SPACING.EXACTLY;f.line_spacing=Pt(28)
    tabs=f.tab_stops;tabs.add_tab_stop(Cm(7.30),WD_TAB_ALIGNMENT.CENTER);tabs.add_tab_stop(Cm(14.55),WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t");p._p.append(_native_equation(kind));p.add_run("\t")
    set_run(p.add_run(f"({number})"),10.5,False,"000000",font="Times New Roman")
    return p

def build():
    doc=Document(str(TEMPLATE));wipe_body(doc)
    # 清理模板遗留作者信息，避免正式提交文件暴露无关编辑元数据。
    doc.core_properties.title=CFG["title"]
    doc.core_properties.subject=CFG["subtitle"]
    doc.core_properties.author=""
    doc.core_properties.last_modified_by=""
    doc.core_properties.comments=""
    # 保留真实模板自身的 A4 页面设置与页边距：上下 2.54 cm，左右 3.175 cm；不增加页码。
    report_title(doc,CFG["title"],CFG["subtitle"])
    # 按模板结构，标题后直接进入正文，不设置独立封面、摘要、关键词、案例编号或摘要表。
    h1(doc,"一、背景、问题、现状")
    h2(doc,"（一）外协力量已成为输电运检的重要延伸")
    para(doc,"随着输电线路规模持续扩大、运检任务不断增加，巡视、隐患排查、现场值守、工单处置和台账采集等大量基础工作需要外协力量参与。外协队伍已经成为输电运检专业的重要执行力量，其任务组织效率和履职质量，直接影响专业管理效能、现场风险发现和任务闭环质量。随着外协参与范围持续扩大，如何保持任务筛选效率和履职真实性，已经成为输电运检专业必须解决的管理问题。")
    para(doc,"一项外协专项任务落下来，管理人员往往先从台账里找线路、找杆塔、核坐标，再把外部风险点一一对应到现场；任务结束后，还要回头核照片、核工单、核处置结果。设备范围越大、专项越多，同一批基础资料就越容易被反复整理。主业人员的时间因此被大量消耗在“查资料、找对象、看照片”这些重复环节上，任务组织和质量核验逐渐成为管理瓶颈。")
    h2(doc,"（二）外协管理面临“工作量大、质量难保证”两大痛点")
    bullet(doc,"工作量大：一项专项排查任务往往需要从多类台账和海量设备中重新寻找目标。以交叉跨越、防鸟、集中燃放点等工作为例，需要反复关联线路、杆塔、坐标和外部信息，再由外协人员逐项查找、现场摸排。任务对象越多、范围越大，人员投入几乎同步增加，大量精力消耗在重复查询和全量排查上。")
    bullet(doc,f"质量难保证：外协履职涉及大量工单、现场反馈和长期积累的历史数据。仅{ALARM_SCOPE}的告警工单数量就超过11万条，同时还沉淀了大量现场反馈照片和历史处置记录。面对数量庞大、时间跨度长、跨工单分散存储的作业数据，管理人员很难依靠人工逐一核验，更难在历史数据中发现同一照片跨时间、跨工单重复使用等异常情况，传统抽查方式难以全面判断外协工作是否真实、规范、到位。")
    bullet(doc,"共同症结：任务侧依赖“人去找”，质量侧依赖“人去查”，管理能力基本随人员投入线性增长，而设备规模和业务数据已经进入规模化增长阶段。单纯增加人员已难以支撑，排查范围难收敛、历史数据难全量核验、业务经验难固化复用等问题日益突出。")
    figure(doc,"pain-points","图1　外协队伍管理面临的两大痛点",12.3)
    page_break(doc)
    # 4 现状量级
    h2(doc,"（三）数据规模决定了“人海作业”不可持续")
    para(doc,f"目前全省约有{CFG['metrics']['province_poles']/10000:.1f}万基输电杆塔、{PROVINCE_LINES_LABEL}。围绕日常运维和专项治理，需要按照不同业务口径持续统计、更新各类台账，例如交叉跨越、防鸟、集中燃放点等，往往都要重新关联线路、杆塔、坐标及外部数据。传统方式主要依靠人工整理、逐项核对，既耗费大量人力，同一基础数据又在不同专项中反复处理，设备规模越大，重复工作越突出。")
    para(doc,"专项排查通常先确认线路和杆号顺序、检查坐标质量，再把外部目标与线路位置对应。交叉跨越、防鸟、集中燃放点的判断规则不同，但线路、杆塔和坐标整理步骤高度重复。")
    para(doc,f"业务过程数据同样持续增长。{ALARM_SCOPE}共形成{CFG['metrics']['alarm_workorders']:,}条告警工单、{CFG['metrics']['alarm_photos']:,}张现场反馈照片；巡视照片每月约{CFG['metrics']['patrol_photos_monthly']/10000:.0f}万张。数据跨线路、跨任务、跨时间积累，人工不仅难以逐一核验，更难与历史记录进行全量比对。以月度{CFG['metrics']['patrol_photos_monthly']/10000:.0f}万张巡视照片为例，若采用简单两两比较，理论组合约77.5万亿对，依靠人工或简单穷举均无法支撑。")
    figure(doc,"province-map","图2　省域线路杆塔任务规模与电压等级分布",12.8)
    para(doc,"40.1万基杆塔、9200余条线路以及持续增长的千万级业务照片，使传统“增加人员、提高频次”的管理方式逐渐触及能力上限。继续增加人手只能有限提高处理量。面向省域输电运检管理，任务筛选和履职监督需要具备批量处理、统一规则和可追溯复核能力，这也是本案例选择外协管理作为切入点的现实原因。")
    page_break(doc)
    # 5 根因
    h2(doc,"（四）根因分析：传统人工方式难以支撑全量管理")
    para(doc,"一是数据“能看不能算”。表格里已经有线路、杆塔、坐标、照片等信息，但来源不同、字段写法不同，杆号顺序和坐标质量也并不一致。人工查看单条记录时问题不大，一旦需要批量计算，相同对象如果名称不统一、坐标异常没有处理，结果就会出现漏算或错算。因此，真正开始计算前，必须先把业务数据整理成能够稳定计算的数据结构。")
    para(doc,"二是业务经验“能做不能复用”。长期从事线路运维的人员知道怎样判断重要跨越、哪些环境容易发生鸟害、什么样的照片值得重点核对，但这些判断往往停留在个人经验和临时表格里。人员变化、任务变化以后，许多步骤又要重新摸索。把经验拆成明确条件并转化为可执行规则，才能在不同线路和不同专项中按同一标准重复执行。")
    para(doc,"三是工具与流程没有完全接上。计算得到一个点位或一对相似照片，只完成了第一步。点位还需要有人去现场确认，照片还需要结合时间、工单和现场细节判断，确认后的结论还要回到任务管理中。计算结果如果停留在单独文件里，业务人员仍要二次整理，实际管理效果会受到很大影响。")
    para(doc,"因此，本案例把重点放在两件事情上：一方面让数字化筛选能力承担适合批量处理的计算和筛选工作，另一方面保留人工对现场真实性、专业风险和处置结果的判断。这样既能扩大覆盖范围，也能保证最后的管理结论由业务人员负责。")
    page_break(doc)
    # 6 模型
    h1(doc,"二、总体思路：构建“增效+提质”的外协数智管理模式")
    h2(doc,"（一）“增效管任务、提质管履职”双线协同")
    para(doc,"针对上述问题，我从输电运检实际业务出发，自主开展数据梳理、规则设计和数字化工具开发，把空间计算用于任务筛选，把图像相似度识别用于照片查重，并归纳为“增效管任务、提质管履职”两条主线：任务侧由系统从全量对象中筛选重点，外协按清单精准核验；履职侧由系统对作业照片全量查重，管理人员重点复核疑似异常照片对。相关算法和数字化工具均由我自主设计、开发和验证，业务筛选流程也由我结合实际管理需求设计并验证。")
    figure(doc,"dual-wheel","图3　外协管理增效提质总体模型",14.4)
    para(doc,"整个模式可以概括为“两条主线、一个机制”：增效管任务、提质管履职；机器负责全量筛选，人员负责精准核验。通过重新划分人机职责，把人的精力从重复劳动转向专业判断。")
    page_break(doc)
    # 7 architecture
    h2(doc,"（二）外协数智管理的共用闭环")
    para(doc,"在双线模式基础上，本案例把计算结果直接接入任务组织、现场执行和主业复核流程。多源数据汇集后，系统按业务规则筛出重点任务并生成清单，外协人员按清单执行，主业复核结果再回写数据和规则，用于后续任务调整。")
    figure(doc,"workflow","图4　外协任务数字化筛选与闭环管理流程",15.2)
    bullet(doc,"多源数据汇集：统一汇集线路、杆塔基础数据，以及交叉跨越相关外部图层、鸟类发生记录、燃放点、气象和节假日等数据。")
    bullet(doc,"空间叠加与规则筛查：把安全距离、风险等级等业务要求转成可计算条件。")
    bullet(doc,"重点任务自动识别：生成高风险、高影响点位清单，减少全量人工排查。")
    bullet(doc,"外协定向巡视：任务和路线精准推送，现场执行并上传反馈。")
    bullet(doc,"主业复核与闭环：确认结果、处置问题并持续优化规则。")
    para(doc,"从实践结果看，本案例的创新主要集中在三个层面。管理层面，把外协管理归纳为“增效管任务、提质管履职”两条主线；方法层面，把空间关系、距离条件和历史照片相似关系转成机器可执行的筛选规则；工程层面，把原有少量特高压照片筛查扩展到全电压等级，支撑千万级巡视照片持续处理。三者共同服务同一目标：机器先处理全量数据，人员集中完成专业核验。")
    para(doc,"上述工作中，我重点承担了基础数据梳理、空间分析方法设计、照片智能筛查方法设计、千万级处理链改造以及业务闭环验证。相关核心算法和数字化工具均由我自主设计、开发和验证，技术成果最终都以任务清单、复核线索和管理结论的形式接回实际业务。")
    para(doc,"空间分析模块负责求交和距离筛选，照片智能筛查模块负责相似照片召回，外协人员负责现场执行，主业人员负责风险判断和结果确认。筛选结果只作为任务线索和复核入口。线路改造、环境变化或现场核验发现规则偏差时，再同步更新基础数据和筛选条件。")
    para(doc,"从技术实现看，这套方法采用分层处理结构。原始业务数据先经过标准化和质量检查，进入统一数据层；随后由规则层把业务判断转成空间关系、距离阈值和图像相似度条件；计算层承担批量求交、距离检索、视觉特征提取和候选召回；输出层只生成需要关注的任务对象或疑似照片对，最后由人工复核并把结果回写。各层输入和输出保持相对稳定，便于在不同专项之间复用。")
    paper_table(doc,"表1　外协数智管理技术架构及输入输出",("层级","主要输入","核心处理","主要输出"),(
        ("数据层","线路、杆塔、坐标、外部地理要素、工单、照片","字段统一、格式转换、完整性检查","可计算的标准数据"),
        ("规则层","业务制度、距离要求、现场经验、复核样本","求交条件、区域条件、距离阈值、相似度阈值","机器可执行规则"),
        ("计算层","标准数据、规则参数","空间计算、特征计算、候选召回、批次调度","候选点位、区段、照片对"),
        ("业务层","计算候选、关联任务信息","任务派发、现场核验、人工终审","业务确认结果"),
        ("回写层","现场结论、复核结论、异常样本","状态更新、样本留存、规则校准","下一轮数据与参数"),
    ),widths=(2.2,4.0,4.5,3.4),note="表中“输出”均为后续业务处理的输入，算法结果不直接替代现场核验和管理定性。")
    para(doc,"分层结构还解决了技术与业务之间的接口问题。例如，交叉跨越和集中燃放点虽然使用不同规则，但都可以调用同一套线路空间数据；告警工单和巡视照片虽然来源不同，也可以共用照片标准化、特征计算和候选复核框架。后续增加新专项时，优先复用已有数据处理和计算模块，只针对业务差异增加规则配置。")
    para(doc,"为了避免不同专项各自维护一套数据，本案例把业务对象按“实体—关系—状态”组织。线路、杆塔、外部点位、工单和照片作为实体；线路—杆塔、工单—照片以及候选照片A—B之间建立关联关系；待核验、确认重复、已闭环等作为状态。数据处理时尽量使用稳定标识关联对象，展示给人员时再恢复线路名称、杆号和业务描述。这样可以降低名称变更、表格列顺序变化对计算链的影响。")
    para(doc,"数据接口设计同时保留原始值和标准化值。例如坐标字段保留原始来源，计算层使用统一格式后的数值；照片记录保留原文件关联，特征层单独保存计算结果。出现异常时既可以看到计算使用了什么数据，也可以回到原始记录核对，避免标准化过程中丢失追溯依据。")
    paper_table(doc,"表2　核心数据对象及关联关系",("对象","稳定关联信息","主要关系","典型状态"),(
        ("线路","线路标识、电压等级","包含多个杆塔和线路段","在运、调整、待核对"),
        ("杆塔","线路标识、杆号（设备标识）","相邻杆塔组成线路段","坐标正常、异常、待复核"),
        ("外部目标","来源标识、几何对象标识","与线路段或杆塔形成空间关系","候选、已核验、已排除"),
        ("工单","工单标识、任务对象","关联一组反馈照片","已完成、待复核、问题闭环"),
        ("照片","照片标识、来源任务","关联特征、候选照片对","正常、候选、已确认"),
    ),widths=(2.2,4.1,4.5,3.6),note="表中强调的是关联原则，具体字段名称随数据源调整；原始值与标准化值均保留追溯关系。")
    page_break(doc)
    # 8 efficiency overview
    h1(doc,"三、具体做法")
    h2(doc,"（一）增效：数字化筛选减少外协无效排查")
    para(doc,"围绕“增效管任务”，最先解决的是交叉跨越排查。交叉跨越空间分析能力跑通后，线路空间数据和计算方法继续用于防鸟、集中燃放点等专项。三个场景的做法基本一致：先从全量对象中筛出重点，再由外协人员按清单现场核验。")
    para(doc,"一是建立可复用的数据底座。对既有杆塔和线路数据进行清洗，检查空坐标、重复点、离群点和杆号顺序，将相邻杆塔连接为线路段；同时对铁路、鸟类活动、燃放点、气象和节假日等外部数据统一标准化，将分散数据整理为可直接计算、可供不同专项重复调用的任务底座。")
    para(doc,"二是把业务经验转化为可计算规则。在统一数据底座基础上，将“是否相交”“是否进入重点活动区域”“是否落入安全距离”等业务判断转化为几何求交、区域叠加和距离阈值，并编写对应计算逻辑，把个人经验固化为可重复执行的筛选规则，批量生成重点核验清单。")
    para(doc,"三是以任务清单组织外协精准核验。候选点位和区段按风险、区域和任务类型整理为清单，定向推送给外协人员并规划巡视路线。外协上传现场结果，主业复核后回写状态，形成“自动筛选—定向执行—反馈复核—规则优化”的闭环。")
    para(doc,"完成线路、杆塔和外部数据整理后，不同专项可以直接调用同一套基础数据和计算规则。外协人员根据自动生成的清单开展现场核验，省去了从全量设备中逐项查找的过程。")
    para(doc,"数据治理是空间计算能够稳定运行的前提。杆塔坐标需要同时满足数值有效、位置合理、线路归属明确和杆号顺序可恢复等条件；外部地理数据则要统一空间参考、几何类型和必要属性。对无法自动判定的异常数据，处理过程中保留原始记录和异常原因，单独进入人工复核，避免异常值直接参与批量计算。")
    paper_table(doc,"表3　空间任务数据预处理与质量控制",("对象","主要质量问题","数字化处理","业务控制"),(
        ("杆塔坐标","缺失、重复、明显离群","完整性校验、重复检测、范围检查","异常记录单独复核"),
        ("线路顺序","杆号缺失、排序错误、跨线路混入","按线路分组并恢复相邻关系","核对线路拓扑与杆号"),
        ("外部线要素","铁路、公路分段过长、属性不统一","几何拆分、属性标准化、范围过滤","保留可追溯来源"),
        ("外部点要素","坐标格式差异、重复点位","坐标统一、去重、范围预筛","核对任务清单来源"),
        ("输出结果","同一对象命中多个规则或多个外部目标","候选归并、关联关系保留","现场核验后确认状态"),
    ),widths=(2.3,4.0,4.2,3.6),note="质量控制优先保证漏项可追溯、异常可复核，避免在数据清洗阶段静默丢弃业务对象。")
    para(doc,"为保证计算结果能够复盘，每条候选还需要保存来源数据、所用规则和关联对象。这样，当现场核验发现误差时，可以判断问题来自原始坐标、业务阈值还是外部数据，并针对原因修正，避免只在最终结果表中人工改数。")
    para(doc,"坐标处理遵循“原始经纬度保留、计算口径明确”的原则。线段求交要求输电线路与外部线状要素处于同一坐标参考；涉及米制距离的业务判断则采用距离换算或投影坐标，避免把经纬度差值直接当作线性距离。输出候选时继续保留原始坐标和业务对象标识，使计算过程与业务台账能够相互对应。")
    para(doc,"线路拓扑构造时，以同一线路内相邻杆塔为基本计算单元。一条线路有n基有效杆塔时，正常情况下形成n−1条相邻线段；如果杆号存在跳号或坐标异常，先记录异常，再决定是否连接。采用档段作为最小单元有两个好处：一是可以把交点、鸟类重点区段等结果精确映射到相邻杆塔之间；二是局部数据修正只影响对应档段，不需要重建整条线路的业务结果。")
    para(doc,"候选结果去重时不能只按坐标取整。多个外部对象可能在相近位置与同一线路相交，或者同一杆塔可能落入多个燃放点影响范围。计算结果因此同时保留业务对象标识和空间关系，先合并完全重复的计算记录，再保留一对多、多对一等真实业务关系，防止为了表格简洁误删需要现场核验的信息。")
    page_break(doc)
    # 9 crossing
    h2(doc,"1.交叉跨越：把逐线排查变成自动求交")
    para(doc,"这套方法最早起于一次交叉跨越排查。早期一次实际数据整理中，需要处理6438个杆塔坐标、249条线路，既有资料中记录了241处铁路跨越结果。按原有方式，需要人员逐条线路、逐个区段翻表和查图，再把铁路、公路位置与杆塔区段对应起来。线路一多，重复查看和漏项就很难完全避免。")
    para(doc,"面对这项任务，我先把相邻杆塔连接成连续线路，再把铁路、公路等外部线状要素放到同一空间关系中计算。空间分析工具采用“包围盒预筛＋精确几何求交”的两级处理：先快速排除明显无关对象，再对保留对象判断是否真正相交。原本依赖人员逐段寻找的问题，由此变成了先计算、后核验。")
    para(doc,"计算前先按线路和杆号顺序把杆塔坐标连接成线段，标记缺失坐标、重复点和明显离群点，再将线路与铁路、公路统一到同一坐标体系。离散杆塔点只有恢复成连续线路段后，才能判断一档导线走廊与外部线状目标之间的空间关系；如果杆号顺序错误，连接出的线段也会偏离真实线路，因此拓扑检查直接影响后续求交可靠性。")
    para(doc,"批量求交采用“包围盒预筛—精确求交”两级结构。实际计算先根据输电线路段的经纬度范围构造最小外包矩形，并向外缓冲0.0005°，利用空间相交关系从铁路、公路等外部线要素中筛出可能相关的对象；只有预筛结果进入后续精确几何求交。精确阶段使用Shapely的intersection计算真实交集，并分别处理点、多点、几何集合以及线状重合等结果。")
    paper_equation(doc,"intersection","1")
    para(doc,"式（1）用于说明有限线段求交的几何原理，其中A、B为输电线路段两个端点，C、D为外部线状要素对应线段的两个端点，t、u为各自线段参数；当两条参数线段存在共同点且参数落在[0,1]区间内时，交点位于两条有限线段上。实际精确求交由Shapely几何运算完成，公式用于解释判定原理。",together=True)
    para(doc,"包围盒预筛的作用是控制候选规模。与线路段外包矩形没有空间关系的外部线要素直接被排除，剩余对象再进入精确求交。0.0005°缓冲用于给预筛范围留出小幅边界余量，降低边界位置因几何范围过紧被提前排除的风险。")
    para(doc,"最终结果保留“输电线路段—外部要素—交点坐标—候选类型”的关联关系。这样现场核验发现某一候选有误时，可以追溯到具体外部要素和原始线路段；同一档线路同时与多条外部线状目标相交时，也不会因为简单去重而丢失业务关系。")
    figure(doc,"crossing","图5　交叉跨越自动筛查及候选分布",15.2)
    para(doc,"在某地市实际应用中，原需十余人、近两周完成的全量排查，空间分析工具数分钟即可完成，同时还补充识别出十余处此前人工排查遗漏的铁路、公路等重要跨越信息。这次排查形成的线路空间数据随后继续用于防鸟、集中燃放点等专项，避免不同专项重复整理线路、杆塔和坐标数据。")
    para(doc,"筛选结果先形成候选清单，由熟悉线路的人员结合图形位置和现场情况复核，确认后的点位再进入正式管理。复核过程中发现的杆号顺序异常、坐标偏移和外部地图要素名称不统一等问题同步修正，为后续防鸟和燃放点筛选提供了更可靠的线路空间数据。")
    para(doc,"算法校核采用“既有结果对照＋新增候选复核”的方式。既有资料中的241处铁路跨越可以作为已知样本，用于检查自动筛查能否覆盖已经掌握的跨越关系；新增候选再由人员逐项查看线路段、铁路位置和现场资料，确认是原资料漏项、外部数据差异还是计算误差。校核重点放在原因定位，不以两个清单数量的简单比较作为结论。")
    para(doc,"边界样本需要单独检查。精确求交结果可能是单个交点、多个交点或线状重合，计算过程按照几何类型分别处理；当交集为线状对象时，以交集中心位置作为后续核验参考。对于无几何对象、空几何和不符合目标类型的外部要素，则跳过并保留原始数据来源，避免异常对象干扰正式候选清单。")
    para(doc,"新增十余处此前人工遗漏的重要跨越后，我又反向检查这些点为什么容易被人工漏掉。常见原因包括线路范围较大、跨越位置不明显、外部要素名称不统一以及需要在多份资料之间切换。统一规则可以对全量线段执行一致的判断条件，降低因人工切换资料和逐段浏览造成的漏项风险，人员随后集中完成少量候选的业务确认。")
    page_break(doc)
    # 10 bird
    h2(doc,"2.防鸟：用空间叠加收敛重点区段")
    para(doc,"交叉跨越空间分析能力形成后，我又把同样的空间计算方法用到防鸟排查中。防鸟分析引入全球生物多样性信息机构（Global Biodiversity Information Facility，GBIF）公开鸟类发生记录，对省域内有效记录进行空间汇总并与输电线路叠加，先筛出鸟类活动相对集中的线路区段；任务复核时再参考湿地、水网等周边生态环境，形成防鸟装置排查和差异化巡视清单。")
    figure(doc,"bird","图6　鸟类活动重点区域与输电线路空间叠加筛查",12.8)
    para(doc,"注：鸟类活动表面使用GBIF.org公开鸟类发生记录计算，线路数据复用报告正式省域线路底图；湿地、水网等环境信息用于任务复核参考，现场风险以实地核验结果为准。",9.5,False,indent=False,line=17,color="444444")
    para(doc,"鸟类活动数据与输电线路叠加后，可以先筛出鸟类活动相对集中的线路区段，再安排外协人员现场核验。这样不需要沿线大范围摸排，防鸟装置排查和差异化巡视的任务范围也更明确。")
    para(doc,"鸟类发生记录存在采集时间和地点分布不均等特点，分析结果只用于缩小排查范围。业务复核时可结合湿地、水网等环境信息，现场再核验鸟巢、防鸟装置状态和周边环境。交叉跨越与防鸟虽然判断规则不同，但都复用同一套线路位置数据和空间计算方法。")
    para(doc,"GBIF原始数据处理采用“坐标有效性检查—省域边界过滤—空间栅格统计—活动强度平滑—重点等级划分”的流程。本次数据共扫描229,927条发生记录，经省界多边形过滤后保留229,561条有效记录，随后落入170×210个经纬度栅格进行数量统计。为降低少量记录极密集区域对整体分级的影响，对网格计数采用ln(1+n)变换后进行邻域平滑。")
    para(doc,"活动等级采用样本分布的分位数划分。对平滑后大于0的有效栅格，分析中分别取70%、84%、94%分位点作为活动区域判定阈值、较活跃分界和高活跃分界，再将活动表面转换到与线路一致的平面坐标进行叠加。分位数规则能够随当前输入记录的空间分布变化，避免长期固定一个记录数量阈值。最终线路筛选仍保留现场核验，GBIF发生记录只负责把大范围线路收敛到更值得关注的区段。")
    paper_equation(doc,"bird_surface","2")
    para(doc,"式（2）表示鸟类活动表面的处理过程，其中nᵢⱼ为第i行、第j列栅格内的有效记录数量，ln(1+nᵢⱼ)用于压缩极高记录密度造成的数量级差异；M₃×₃表示一次3×3邻域均值平滑，上标7表示重复应用该算子，即连续执行7次3×3邻域均值平滑。平滑后的有效栅格再按样本分位数划分活动等级，降低少量记录极密集区域对全省分级结果的过度影响。该活动表面只用于空间筛选，不作为鸟害概率或设备故障概率。",together=True)
    para(doc,"线路与活动表面叠加时，以线路段位置读取对应活动等级，达到重点条件的区段进入候选集合；同一线路可能有多个区段连续命中，输出时再按线路和相邻区段归并，形成便于外协执行的核验范围。现场反馈若显示某些区域长期没有鸟类活动或防鸟设施状态稳定，可作为后续规则调整的业务依据。")
    page_break(doc)
    # 11 fireworks
    h2(doc,"3.集中燃放点：复用空间底座快速出清单")
    para(doc,"2025年底接到集中燃放点周边输电杆塔排查任务后，我直接复用前期形成的线路空间数据和距离计算方法，只调整距离参数和筛选规则，算上修改调试不到半小时即完成批量计算和清单输出。")
    para(doc,"这一场景的核心计算是“燃放点—杆塔”的批量距离关系。每个燃放点作为外部点目标进入统一空间计算流程，业务距离作为可配置参数；计算先用经纬度范围做低成本预筛，再计算燃放点与杆塔之间的实际距离。命中杆塔同时保留所属线路标识，由杆塔结果汇总出涉及线路，避免后续人员再从地图位置反查设备。")
    para(doc,"多个燃放点同时计算时，一基杆塔可能命中多个点位，同一线路也可能有多基杆塔进入影响范围。计算结果先保留全部命中关系，再按杆塔和线路归并形成核验清单，并保留对应燃放点来源。这样既能控制清单重复，又不会丢失“一基杆塔对应多个外部风险源”的信息。")
    para(doc,"距离参数与计算主体分离后，新任务只需调整业务参数即可复用。参数改变会重新生成候选范围，杆塔坐标清洗、低成本范围预筛、结果归并和现场核验字段都可以沿用前期模块，这也是该任务能够在较短时间内完成修改调试和清单输出的技术原因。")
    figure(doc,"fireworks","图7　集中燃放点影响范围与周边杆塔筛查",12.8)
    para(doc,"一次批量计算即可给出集中燃放点周边需要核验的杆塔和线路清单，外协人员按清单到现场确认，不再逐个燃放点、逐基杆塔查询判断。交叉跨越、防鸟和集中燃放点虽然业务对象不同，但实际都采用“机器先筛、人员核验”的办法。")
    para(doc,"三个空间专项可以用“点—线—面”三个空间对象统一理解。交叉跨越对应“线”，重点判断输电线路与铁路、公路等线状要素是否相交；防鸟对应“面”，先将鸟类发生记录聚合为活动重点区域，再与输电线路叠加筛选重点区段；集中燃放点对应“点”，以燃放位置为中心，按业务距离筛选周边杆塔并汇总涉及线路。三个场景共同覆盖点、线、面三类空间对象，并共用同一套线路杆塔空间底座和现场核验流程。")
    paper_table(doc,"表4　三类空间专项的点—线—面对象与筛选逻辑",("专项","空间对象","核心筛选逻辑","人员侧核验重点"),(
        ("集中燃放点","点","以燃放点为中心按业务距离筛选周边杆塔","点位有效性、设备实际受影响情况"),
        ("交叉跨越","线","线路段与铁路、公路等线状要素求交","跨越类型、现场位置、管理属性"),
        ("防鸟","面","鸟类活动记录聚合成重点区域，与线路叠加筛选","鸟巢、防鸟装置、周边生态环境"),
    ),widths=(2.5,2.5,5.2,4.1),note="三类专项分别以点、线、面为分析入口，共用线路与杆塔空间底座，业务差异主要体现在筛选规则和核验重点。")
    page_break(doc)
    # 12 quality problem
    h2(doc,"（二）提质：照片查重强化外协履职质量监督")
    para(doc,f"任务侧最直观的问题是“去哪查”，质量侧更难的是“这次到底有没有查到位”。外协人员完成任务后会留下反馈照片，但管理人员面对的是跨工单、跨时间不断累积的历史画面。一张照片单独看很正常，只有把它和过去几周、几个月的照片放在一起，才可能看出重复使用。{ALARM_SCOPE}共形成{CFG['metrics']['alarm_workorders']:,}条告警工单和{CFG['metrics']['alarm_photos']:,}张现场反馈照片，人工记忆和逐张比对已经无法覆盖这样的规模。")
    figure(doc,"photo-scale","图8　告警工单照片全量筛选与人工复核结果",15.2)
    para(doc,f"系统对{CFG['metrics']['alarm_photos']:,}张现场反馈照片进行全量筛选，形成{CFG['metrics']['alarm_candidates']:,}对疑似相似照片，其中{CFG['metrics']['alarm_reviewed']:,}对已完成人工复核，确认重复{CFG['metrics']['alarm_confirmed_pairs']:,}对、确认不同{CFG['metrics']['alarm_confirmed_different']:,}对，另有{CFG['metrics']['alarm_pending']:,}对待复核。管理人员不再需要在11万余张照片中逐张查找，只需集中复核筛选出的疑似照片对。目前，告警工单反馈照片查重和巡视照片查重均已在省公司层面开展试点，并用于生产管控中心运检工作质量远程督查。")
    para(doc,"重复照片既有直接复用，也有修改水印、裁剪或调色后再次提交；同时还存在画面相似但车辆位置、吊臂姿态确有变化的正常照片。照片智能筛查模块只负责把疑似照片对找出来，复核人员再查看两张原图、拍摄时间、线路、工单对象和现场细节，确认照片是否真实对应当次任务。")
    page_break(doc)
    # 13 innovation 1
    h2(doc,"1.告警工单照片查重：从零建立全量筛查方法")
    para(doc,"当时告警工单反馈照片没有可直接使用的全量查重手段。我先确定什么样的照片应进入疑似重复候选，再开发照片智能筛查模块，设计机器筛选、原图对照和人工终审流程，并用真实业务样本反复校准参数。由此建立了告警工单照片全量查重方法。")
    figure(doc,"photo-cases","图9　真实复核案例：同图修改水印日期与高相似但不同照片",12.2)
    para(doc,"图9上部三张照片主体、构图和现场细节一致，但水印日期分别显示为06-27、06-10和05-30，人工复核后确认属于同图修改水印日期后重复使用；下部两张照片虽然相似度达到97.63%，但车辆位置、吊臂姿态和现场物体存在真实变化，人工确认并非重复。全量查重先把疑似问题筛出来，再由管理人员结合工单和现场情况判断。")
    para(doc,"省公司试点中，照片智能筛查模块筛出一组500千伏特殊通道高相似照片，管理人员随即调取原始巡视记录逐项核对，最终确认同一区段两次人工巡视使用了重复照片。由于该记录无法真实反映当时通道情况，相关问题被认定为人工巡视不到位，并定性为较大运检质量问题。")
    para(doc,"目前，照片查重已在省公司生产管控中心实际应用，并累计形成照片重复类问题19项，其中告警工单反馈照片重复11项、人工巡视照片重复8项。")
    page_break(doc)
    # 14 algorithm
    h2(doc,"2.照片查重：pHash先筛，CLIP再比")
    para(doc,"照片查重的思路可以概括为两道筛选。第一道先求“快”：利用感知哈希（perceptual hash，pHash）和汉明距离，从海量照片中迅速找出整体构图接近的组合；第二道再求“准”：利用对比语言—图像预训练模型（Contrastive Language–Image Pre-training，CLIP）计算图像向量相似度，对经过裁剪、调色、旋转或局部遮挡后仍可能表达同一现场的照片继续复核。现有告警工单规则采用“pHash汉明距离＜10且CLIP相似度＞0.80”生成候选。")
    para(doc,"阈值确定时同时考虑漏报和人工复核量。阈值过严可能漏掉裁剪、压缩或修改水印后的重复照片，过宽又会产生大量正常相似候选。pHash负责低成本筛出整体结构接近的照片，CLIP再对这些候选做语义比较，参数通过真实工单样本反复校准，并把候选量控制在人工能够复核的范围内。")
    para(doc,"pHash的输入是经过统一方向和尺寸处理后的图像。第一阶段提取低频结构信息并生成固定长度的感知指纹，两张照片的指纹通过汉明距离比较，距离越小表示整体构图越接近。与普通文件哈希相比，这种方法对重新压缩、格式变化、轻度亮度变化和局部水印修改更稳定，适合承担海量照片的第一轮快速召回。")
    paper_equation(doc,"hamming","3")
    para(doc,"式（3）表示两个感知指纹在全部比特位置上的差异数量，其中h₁、h₂为两张照片的感知指纹，L为指纹长度，I(·)为指示函数，条件成立取1，否则取0。当前告警工单规则要求pHash汉明距离小于10，意味着只有整体结构足够接近的照片才进入下一层。pHash阶段重点追求低成本召回，因此它只判断“值得继续比较”，不会直接对照片重复性质作结论。",together=True)
    para(doc,"CLIP阶段将候选照片映射为图像特征向量，再计算向量相似程度，用于补充pHash主要关注整体构图的局限。实际样本中曾出现CLIP相似度达到97.63%、最终人工确认不同的照片，说明高相似度仍需要结合车辆位置、吊臂姿态、作业对象和时间关系进行业务复核。")
    para(doc,"双阈值采用同时满足的方式生成候选：结构层先控制构图差异，语义层再控制内容差异。两个阈值承担的角色不同，pHash阈值主要影响第一层召回范围，CLIP阈值影响进入人工队列的候选纯度。参数调整时需要观察确认重复、确认不同和待复核三类结果的变化，防止只追求候选数量下降而增加漏检风险。")
    paper_table(doc,"表5　照片查重分层筛选流程及责任边界",("阶段","输入","主要处理","输出及用途"),(
        ("照片标准化","原始反馈照片、巡视照片","方向修正、尺寸与颜色通道统一、异常文件检查","统一图像输入"),
        ("pHash召回","标准化图像","感知指纹计算、汉明距离筛选","结构近似候选"),
        ("CLIP复核","pHash候选","图像特征向量提取、相似度计算","疑似照片对"),
        ("候选归并","多批次候选","照片对标准化、重复关系合并、关联任务信息","人工复核队列"),
        ("人工终审","原图、相似度、任务和时间信息","核对现场变化、水印、对象和业务背景","确认重复、确认不同、待复核"),
    ),widths=(2.3,3.5,4.6,3.8),note="当前告警工单候选规则采用pHash汉明距离＜10且CLIP相似度＞0.80；阈值用于候选生成，不直接形成质量问题定性。")
    para(doc,"111,519张反馈照片经过分层筛选后形成5,472对疑似相似照片，已有4,630对完成人工复核，其中确认重复348对、确认不同4,282对，另有842对待复核。若不附加任何业务约束，这批照片的理论两两组合约62.18亿对；这一数量级只用于说明原始搜索空间，实际处理通过候选生成把人工复核和高成本计算集中到少量照片对上。348对确认重复占已复核候选的7.52%，这一比例反映候选队列的复核结果，不代表全部照片的重复率。")
    para(doc,"4,282对“确认不同”覆盖了同位置、同设备、构图接近但现场确有变化的正常场景，可与确认重复样本一起用于检查阈值是否偏向某一类画面。水印修改、裁剪、压缩和高相似不同等典型正反样本持续保留，使参数调整既关注重复照片能否被召回，也关注人工复核量是否保持在合理范围。")
    para(doc,"评价方法分为三个层次：先看照片是否完整进入标准化和特征计算，再看候选是否具有复核价值，最后看候选能否通过原图、任务记录和现场信息形成可靠业务结论。由于目前没有覆盖全量照片两两关系的人工真值集，本案例不使用缺乏完整验证基础的准确率、召回率等指标，而采用候选量、复核量、确认结果和实际通报问题量描述运行成效。")
    paper_table(doc,"表6　告警工单照片筛查规模与评价口径",("指标","数值","技术含义","使用注意"),(
        ("反馈照片总量","111,519张","全量筛查的数据入口","表示照片数量，不表示组合数量"),
        ("无约束理论两两组合","约62.18亿对","说明未附加业务约束时的搜索空间规模","仅用于量级分析"),
        ("机器候选","5,472对","双阶段规则输出的人工复核入口","候选不等同于重复"),
        ("已人工复核","4,630对","已有明确业务判断的候选","另有842对待复核"),
        ("确认重复","348对","人工终审确认的重复照片关系","属于业务结论"),
        ("已复核候选重复占比","7.52%","348÷4,630，反映已复核候选结果","不能解释为全量照片重复率"),
    ),widths=(3.1,2.5,4.8,4.0),note="理论组合用于说明搜索空间规模；实际处理通过分层候选生成避免进行全量高成本精细比对。")
    para(doc,"表5已将照片标准化、pHash召回、CLIP复核、候选归并和人工终审串成完整链路。业务应用时，管理人员直接查看带有相似度、任务对象和时间信息的疑似照片对，复核人员重点核对现场变化、水印和业务背景，再给出确认重复、确认不同或待复核结论。")
    page_break(doc)
    # 15 innovation 2
    h2(doc,"3.巡视照片：把查重范围扩展到全电压等级")
    para(doc,"告警工单查重跑通后，我开始处理规模更大的巡视照片。此前由电力信息公司提供的既有照片查重能力主要覆盖少量特高压巡视照片，尚无法满足全电压等级海量照片的常态化筛查需求。扩展到全电压等级后，每月需要处理约1245万张巡视照片；若不附加任何约束，理论两两组合约77.5万亿对，原有处理方式无法直接放大。因此重新设计了候选生成、特征复用、分批调度、结果去重和断点续算等处理流程。")
    para(doc,"与原有处理范围相比，主要有三点变化：",14,True,keep=True)
    bullet(doc,"覆盖范围：从少量特高压扩展到全电压等级。")
    bullet(doc,"处理规模：从小规模照片筛查提升至月度千万级。")
    bullet(doc,"运行方式：从单次处理转为可分批、可续算、可持续运行的规模化全量筛查。")
    para(doc,"完成上述改造后，查重范围从少量特高压照片扩展到全电压等级，目前能够按月处理约1245万张巡视照片。")
    page_break(doc)
    # 16 scale technical
    h2(doc,"4.千万级处理：让1245万张照片稳定跑完")
    para(doc,"当处理对象从11万级反馈照片扩大到覆盖全电压等级的月度千万级巡视照片时，难点也随之变化：算法能够找出相似照片只是起点，整个月的数据能否稳定跑完才决定这项能力能不能长期使用。为此，计算路径按“特征预计算—低成本候选召回—高成本语义复核”分层，把真正需要精细比较的数据压缩到较小范围。")
    para(doc,"1245万张照片若无约束两两组合，理论规模约77.5万亿对。面对这样的数量级，单纯提高一次比较的速度意义有限，更重要的是尽早缩小比较范围。处理链先为每张照片计算可复用特征，再用低成本规则生成候选，只有这些候选进入后续语义比较；已经完成的特征直接复用，避免同一张照片反复计算。")
    paper_equation(doc,"combinations","4")
    para(doc,"式（4）中N为参与查重的照片数量，C(N,2)为两两组合数量，说明全组合规模随照片数量近似按平方增长。N达到千万级以后，即使每次比较耗时很短，组合总量仍会迅速超过常规任务能够承受的范围。因此工程设计把主要精力放在候选空间收敛、特征复用和任务可恢复性上，避免把优化重点局限在单次推理速度。",together=True)
    para(doc,"千万级数据还带来运行可靠性问题。批处理引擎把任务拆成稳定批次，每批结束后记录进度和结果；机器重启、网络中断或单批数据异常时，只重跑受影响批次，已经完成的进度和特征继续复用，不同批次结果统一去重归并。")
    para(doc,"工程实现中把“照片特征”和“照片两两关系”分开管理。单张照片的特征只需要计算一次，后续不同批次、不同时间范围检索时可以直接复用；候选照片对则采用稳定的两端标识排序，保证A-B与B-A不会形成两条重复结果。每个批次同时记录输入范围、完成状态和输出数量，为断点续算和异常追踪提供依据。")
    para(doc,"批次设计还需要控制内存和中间结果规模。读取照片、提取特征、召回候选和语义复核分别按可控数据块执行，完成一批即释放无关中间对象并保存结果。某一批次失败时，调度层根据状态记录重新执行该批次，其余已完成数据保持有效。这样可以把长时间运行任务拆成可检查、可恢复的执行单元。")
    paper_table(doc,"表7　千万级照片处理的工程化机制",("机制","主要问题","处理方式","工程作用"),(
        ("特征预计算","同一照片被多次读取和推理","单图特征计算后持久复用","降低重复计算量"),
        ("低成本召回","全组合数量约77.5万亿对","先用低成本特征缩小候选范围","减少高成本比较"),
        ("分批调度","单任务数据量过大","按批次读取、计算、保存并释放资源","控制内存与运行窗口"),
        ("断点续算","长任务可能被重启或异常中断","记录已完成批次和处理位置","避免整月数据从头运行"),
        ("异常重试","个别文件或批次处理失败","隔离失败批次并单独重跑","提高整体任务完成率"),
        ("结果归并","跨批次产生重复照片关系","统一照片对顺序并去重汇总","生成唯一复核清单"),
    ),widths=(2.2,3.9,4.6,3.7),note="工程目标是在保持结果可追溯的前提下，使全电压等级月度照片筛查能够持续运行。")
    para(doc,"为了让批处理结果具备可审计性，每个执行批次至少需要记录数据范围、任务状态、已完成位置、异常数量和输出候选数量；照片特征记录与原始照片保持稳定关联，候选记录则保存两端照片标识、pHash汉明距离、CLIP相似度和关联业务对象。这样在复核阶段可以从候选回到原图，在异常排查阶段也可以从原图追到特征和批次。")
    paper_table(doc,"表8　千万级查重关键记录与复现信息",("记录对象","关键字段及信息","主要用途"),(
        ("照片记录","照片标识、来源任务、拍摄时间、文件状态","关联原图与业务对象，识别异常文件"),
        ("特征记录","照片标识、pHash、CLIP特征状态、计算版本","复用已完成特征，支持版本核对"),
        ("规则记录","pHash汉明距离阈值、CLIP相似度阈值、规则版本","解释候选生成条件，支持参数复盘"),
        ("批次记录","输入范围、开始状态、结束状态、完成位置、异常计数","断点续算、失败隔离和运行审计"),
        ("候选记录","照片A、B标识、pHash汉明距离、CLIP相似度、批次来源","人工复核、候选去重和结果复现"),
        ("复核记录","确认重复、确认不同、待复核、复核备注、关联问题","回写管理结论并形成后续校准样本"),
    ),widths=(2.6,6.9,4.9),note="记录应支持从管理结论反向追溯到候选、规则、特征、原始照片和执行批次，并可复现当次候选生成条件。")
    para(doc,"运行监测保留三类关键信号：任务是否持续推进、异常数据是否被记录、候选数量是否突然发生数量级变化。验证样本同时覆盖完全重复、水印修改、裁剪、压缩、高相似但现场不同以及异常文件，既检查重复照片能否进入候选，也检查正常变化能否在人工终审中得到区分。")
    para(doc,"长周期任务还要保证重复运行结果稳定。同一批数据在规则和特征版本不变时，候选集合应保持一致；任务中断后从已保存状态继续，已完成批次不重复产生复核记录。月度数据或规则发生变化时，则按变化范围处理：新增或修改照片只更新相关特征，调整双阈值可以直接复用已有特征，更换特征模型后再重算受影响部分。结合表8保存的输入范围、特征版本、规则参数和批次状态，可以复现当次候选生成条件。")
    para(doc,"照片智能筛查模块只负责给出疑似照片对，最终是否属于重复使用仍由管理人员结合任务时间、作业对象、现场细节和业务要求判断。这样既能覆盖海量历史照片，又不会把算法相似度直接当作管理结论。")
    h2(doc,"5.技术边界与误差控制：让结果可解释、可复核")
    para(doc,"数字化筛选扩大了管理覆盖范围，同时也带来新的误差来源。空间计算依赖坐标和线路拓扑准确性，公开生态数据存在采样偏差，图像相似度存在误报与漏报。系统设计中保留异常数据、候选依据和人工终审环节，使每类误差都能在对应环节被发现和修正。")
    paper_table(doc,"表9　主要技术边界及控制措施",("技术环节","主要边界或误差来源","控制措施","最终责任"),(
        ("线路空间计算","坐标偏移、杆号顺序异常、外部要素误差","异常检测、拓扑核对、候选现场复核","主业人员确认跨越关系"),
        ("鸟类活动分析","公开发生记录时空分布不均，不能直接代表实时风险","仅作任务范围参考；业务复核结合周边环境","现场核验确认鸟害风险"),
        ("照片相似度筛选","正常场景高度相似、裁剪遮挡可能影响召回","双阶段筛选、正反样本校准、保留人工终审","管理人员认定照片是否重复"),
        ("千万级批处理","中断、异常文件、跨批次重复候选","批次状态记录、断点续算、失败隔离、结果去重","系统核查完整性，人员确认业务结论"),
    ),widths=(2.5,4.1,4.0,3.7),note="技术边界在设计阶段明确标注，算法输出始终作为任务线索和证据候选。")
    para(doc,"从应用角度看，技术可靠性由“数据质量、算法筛选、人工核验”三部分共同保障。空间任务中重点关注漏掉应核验对象的风险，照片任务中重点控制漏报风险和人工复核负荷。后续每次现场确认和照片终审都会产生新的正、反样本，持续用于检查规则是否偏离实际业务。")
    page_break(doc)
    # 17 effects
    h1(doc,"四、实施效果")
    h2(doc,"（一）增效：从“大范围找目标”到“拿清单去核验”")
    para(doc,"交叉跨越排查最能体现任务组织方式的变化。过去拿到的是一项大范围排查要求，人员需要在多份资料和地图之间来回切换，逐条线路、逐个区段寻找目标；采用数字化筛选后，重点点位和区段先形成清单，外协人员直接围绕清单到现场确认。某地市重要跨越排查原需十余人、近两周完成，空间分析工具数分钟即可完成第一轮全量计算，同时补充发现十余处此前人工遗漏的重要跨越信息。")
    para(doc,"这次排查形成的线路空间底座随后继续用于防鸟和集中燃放点。2025年底集中燃放点任务到来后，仅需调整距离参数和筛选规则，算上修改调试不到半小时即可形成周边杆塔和涉及线路清单。同一批基础数据由一次性整理转为多专项复用，主业人员减少反复查资料，外协人员把更多时间留给现场核验。")
    h2(doc,"（二）提质：从“抽到才发现”到“全量先筛一遍”")
    para(doc,f"告警工单照片筛查从{CFG['metrics']['alarm_photos']:,}张反馈照片起步，形成{CFG['metrics']['alarm_candidates']:,}对疑似相似照片，目前已确认重复{CFG['metrics']['alarm_confirmed_pairs']:,}对。随后，照片查重能力扩展到巡视业务和全电压等级，省公司试点累计形成照片重复类问题{CFG['metrics']['province_trial_duplicate_issues']}项，其中告警工单反馈照片重复{CFG['metrics']['province_trial_alarm_duplicate_issues']}项、人工巡视照片重复{CFG['metrics']['province_trial_patrol_duplicate_issues']}项。")
    para(doc,"其中，500千伏特殊通道问题由机器筛出高相似线索，管理人员调取两次巡视原始记录逐项比对，确认同一区段重复使用照片，随后认定为人工巡视不到位，并定性为较大运检质量问题。该结果已直接用于履职质量判断和管理闭环。")
    h2(doc,"（三）管理方式：机器处理全量，人员负责判断")
    para(doc,"交叉跨越、防鸟和集中燃放点共用线路空间底座，告警工单和巡视照片共用照片智能筛查方法。机器处理规模大、重复性高的筛选工作，外协人员和主业人员集中完成现场确认、业务终审和问题处置；照片查重已在省公司生产管控中心实际应用。")
    para(doc,"推广不依赖新增现场硬件，主要复用现有线路、杆塔、工单和照片数据。其他地市完成字段映射、坐标统一、规则参数配置和复核流程衔接后，即可复用现有方法。复制路径可以概括为三件事：接数据、配规则、走闭环。")
    paper_table(doc,"表10　案例实施成效与管理方式变化",("管理环节","传统方式","数智协同方式","实际成效"),(
        ("专项任务筛选","人员从全量对象中逐项寻找","空间分析先形成重点点位和区段清单","原需十余人、近两周完成的全量排查，第一轮全量计算缩短至数分钟，并补充发现十余处遗漏"),
        ("专项能力复用","不同专项重复整理线路与坐标","共用线路空间底座，仅调整业务规则","集中燃放点任务修改调试不到半小时即可形成核验清单"),
        ("外协履职监督","依赖有限抽查和人员历史记忆","照片全量筛查后集中复核异常候选","省公司试点形成19项照片重复问题，其中1项500千伏特殊通道问题定性为较大运检质量问题"),
        ("照片处理范围","既有能力主要覆盖少量特高压照片","扩展到全电压等级并采用批处理机制","目前能够按月处理约1245万张巡视照片"),
    ),widths=(2.5,3.7,4.1,4.1),note="成效既包括处理效率，也包括任务组织方式、履职监督范围和可复制能力的变化。")
    h2(doc,"（四）推广价值：从单项成果到可复制机制")
    para(doc,"应用范围由交叉跨越逐步扩展到防鸟、集中燃放点和照片查重，其中照片查重已在省公司生产管控中心实际应用。空间专项复用线路空间底座，照片查重复用统一筛查流程。")
    para(doc,"可复制的核心可以归纳为三层。第一层是数据底座，把线路、杆塔、工单和照片整理成稳定可关联的数据；第二层是业务规则，把求交、区域叠加、距离条件和照片相似关系转成机器可执行条件；第三层是闭环流程，把筛选结果交给外协现场核验或管理人员终审，再将结论回写。专项对象和参数可以变化，这三层结构保持稳定，因此具备跨地市、跨业务复制的基础。")
    para(doc,"其他地市落地时可以从一个典型专项起步：先完成字段映射和坐标统一，再配置本地业务规则，最后接入任务派发、现场核验和结果回写。跑通一条闭环后，再逐步扩展到其他专项。这样既降低一次性改造范围，也能在真实业务中同步校正规则和数据质量，使“接数据、配规则、走闭环”成为可直接执行的推广路径。")
    figure(doc,"outcomes","图10　外协管理由粗放管理向数智协同管理转型",12.8)
    # 18 recommendations
    h1(doc,"五、推广应用建议")
    h2(doc,"1.建立外协任务数字筛选机制")
    para(doc,"统一坐标、线路顺序和任务关联字段，以交叉跨越、防鸟、燃放点等专项为模板，把全量对象自动筛选为候选清单，形成“系统先筛、外协精准核验”的常态化任务组织方式。")
    h2(doc,"2.建立外协质量全量监督机制")
    para(doc,"将照片查重作为外协履职质量监督的常态化手段，按任务、时间和对象形成异常候选清单，并由管理人员重点复核、闭环处置。")
    h2(doc,"3.坚持机器全量筛选、人员精准核验")
    para(doc,"算法与计算模块追求全量覆盖和高效召回，人员负责业务真实性判断和处置闭环，避免把算法相似度直接等同于管理结论。")
    h2(doc,"4.建立数据持续更新与成效评价机制")
    para(doc,"下一步结合省公司试点情况，把现场核验结果、设备变更信息和照片复核结论及时回写，并同步记录专项排查耗时、候选核验量和异常发现数量。推广前重点统一数据接口、规则参数和处理流程，再向其他地市公司推广应用。")
    para(doc,"推广时要明确线路、杆塔、坐标和外部数据的更新责任，并固定照片候选生成、人工复核和问题反馈流程。评价直接看实际工作结果：任务侧统计候选数量、现场核验量和人工耗时，质量侧统计候选照片对、确认问题数量和复核工作量，再根据结果调整规则。")
    h2(doc,"结语")
    para(doc,"本案例最直接的成效体现在两个方面。任务侧，一次原需十余人、近两周完成的交叉跨越排查，第一轮全量计算现在只需数分钟，同时补充发现十余处人工漏项；质量侧，照片查重从告警工单起步，随后扩展到巡视照片，并在省公司试点中发现500千伏特殊通道重复用图这一较大运检质量问题。两个场景分别解决了“任务太多找不过来”和“照片太多查不过来”的现实矛盾。",14,False,together=True)
    para(doc,"由此形成的做法可以归纳为一句话：增效管任务、提质管履职。数字化工具负责处理规模大、重复性高的数据，外协和主业人员把精力放到现场确认、风险判断和问题处置。推广时按照“接数据、配规则、走闭环”的顺序完成本地化配置，即可在现有业务数据和管理流程上复用这套方法。",14,False,together=True)
    OUT.parent.mkdir(parents=True,exist_ok=True);doc.save(OUT)
    print(f"[report] generated {OUT}")

if __name__=="__main__":build()
